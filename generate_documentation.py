from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_documentation():
    doc = Document()
    
    # Configuration des styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Inter'
    font.size = Pt(11)

    # PAGE DE GARDE
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('\n\n\nSYSTEME DE GESTION DES VISITEURS (VMS)\n')
    run.bold = True
    font = run.font
    font.size = Pt(28)
    font.color.rgb = RGBColor(37, 99, 235) # Bleu primaire

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run('DOCUMENTATION TECHNIQUE ET FONCTIONNELLE COMPLETE')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(31, 41, 55)

    doc.add_page_break()

    # --- SECTION 1: DOCUMENTATION TECHNIQUE ---
    doc.add_heading('1. Documentation Technique', level=1)
    
    doc.add_heading('1.1. Introduction et Objectifs', level=2)
    doc.add_paragraph(
        "Ce projet a été conçu pour répondre aux besoins critiques de sécurité et de traçabilité des accès physiques. "
        "Le système permet d'identifier chaque visiteur, d'automatiser la saisie des données via l'OCR et de conserver "
        "un historique inaltérable des mouvements."
    )

    doc.add_heading('1.2. Architecture logicielle', level=2)
    doc.add_paragraph(
        "L'application repose sur une architecture Django 6.0 modernisée, utilisant une base de données relationnelle "
        "et des bibliothèques de pointe pour le traitement des documents."
    )
    
    tech_table = doc.add_table(rows=1, cols=2)
    tech_table.style = 'Table Grid'
    hdr_cells = tech_table.rows[0].cells
    hdr_cells[0].text = 'Composant'
    hdr_cells[1].text = 'Détails'
    
    technologies = [
        ('Framework Backend', 'Django 6.0.3 (MVT Architecture)'),
        ('Reconnaissance de texte', 'Tesseract OCR (Extraction CNIB)'),
        ('Moteur de Rendu PDF', 'xhtml2pdf / ReportLab'),
        ('Interface Frontend', 'HTML5 / CSS3 (Bootstrap 5.3 + Custom Premium Styling)'),
        ('Base de Données', 'SQLite (Structure SQL optimisée)'),
        ('Sécurité', 'Authentification Django Session-based + CSRF Protection'),
    ]
    
    for comp, tech in technologies:
        row_cells = tech_table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = tech

    doc.add_heading('1.3. Modèles de Données (Schema)', level=2)
    doc.add_paragraph("Nous avons modélisé le système autour de 5 entités centrales :")
    
    entities = [
        ('Porte', 'Représente un point d\'accès. Chaque visite est rattachée à une porte.'),
        ('Visiteur', 'Stocke les données d\'identité, photos et scans (CNI recto/verso).'),
        ('Service', 'Le département de destination au sein de l\'institution.'),
        ('Visite', 'Lien entre un visiteur, un service et une porte (avec horodatage précis).'),
        ('LogAction', 'Trace technique de chaque modification effectuée par un administrateur.'),
    ]
    
    for name, desc in entities:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name} : ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # --- SECTION 2: DOCUMENTATION FONCTIONNELLE ---
    doc.add_heading('2. Documentation Fonctionnelle', level=1)
    
    # 2.1 Authentification
    doc.add_heading('2.1. Mire de Connexion', level=2)
    doc.add_paragraph(
        "L'accès est protégé. Chaque utilisateur (Admin ou Agent) doit s'authentifier. "
        "L'interface de connexion est épurée et sécurisée."
    )
    
    img_login = r'C:\Users\HP envy\.gemini\antigravity\brain\12ffd404-a958-4063-b6c3-9638bfe1ecab\vms_login_mockup_1776166634244.png'
    if os.path.exists(img_login):
        doc.add_picture(img_login, width=Inches(6))
        p = doc.add_paragraph('Figure 1 : Interface de connexion sécurisée')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2.2 Dashboard
    doc.add_heading('2.2. Tableau de Bord Intuitif', level=2)
    doc.add_paragraph(
        "Dès la connexion, nous accédons au tableau de bord. Il présente les indicateurs clés : "
        "nombre de visiteurs sur place, statistiques du jour, et graphiques de tendance."
    )
    
    img_dash = r'C:\Users\HP envy\.gemini\antigravity\brain\12ffd404-a958-4063-b6c3-9638bfe1ecab\vms_dashboard_mockup_1776166504386.png'
    if os.path.exists(img_dash):
        doc.add_picture(img_dash, width=Inches(6))
        p = doc.add_paragraph('Figure 2 : Tableau de bord et statistiques en temps réel')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 2.3 Gestion Visiteurs
    doc.add_heading('2.3. Gestion des Visiteurs et Liste', level=2)
    doc.add_paragraph(
        "Nous pouvons consulter la liste de tous les visiteurs et effectuer des recherches par nom ou numéro CNIB. "
        "Chaque fiche visiteur contient l'historique complet de ses passages."
    )
    
    img_list = r'C:\Users\HP envy\.gemini\antigravity\brain\12ffd404-a958-4063-b6c3-9638bfe1ecab\vms_visitor_list_mockup_1776166742693.png'
    if os.path.exists(img_list):
        doc.add_picture(img_list, width=Inches(6))
        p = doc.add_paragraph('Figure 3 : Liste des visiteurs et recherche multicritères')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2.4 OCR
    doc.add_heading('2.4. Enregistrement par Scan OCR', level=2)
    doc.add_paragraph(
        "Pour gagner du temps, nous utilisons le module OCR. En téléchargeant une photo de la CNIB, "
        "le système remplit automatiquement le formulaire, évitant ainsi les erreurs de saisie."
    )
    
    img_ocr = r'C:\Users\HP envy\.gemini\antigravity\brain\12ffd404-a958-4063-b6c3-9638bfe1ecab\vms_ocr_scan_mockup_1776166761168.png'
    if os.path.exists(img_ocr):
        doc.add_picture(img_ocr, width=Inches(6))
        p = doc.add_paragraph('Figure 4 : Module de reconnaissance de texte (OCR)')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('2.5. Gestion des Entrées/Sorties', level=2)
    doc.add_paragraph(
        "Lorsqu'un visiteur arrive, nous créons une 'Visite' rattachée à un service. "
        "Lors de son départ, il suffit de cliquer sur 'Sortie' pour libérer la place."
    )

    doc.add_heading('2.6. Rapports et Audit', level=2)
    doc.add_paragraph(
        "L'administrateur peut générer des rapports PDF pour n'importe quelle période. "
        "Le journal d'actions (Logs) permet de savoir exactement qui a fait quoi et quand."
    )

    filename = 'Documentation_VMS_Finale.docx'
    doc.save(filename)
    print(f"Documentation générée avec succès : {filename}")

if __name__ == "__main__":
    create_documentation()
